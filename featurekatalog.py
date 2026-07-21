#!/usr/bin/env python3
import re
import sys

import docx
import yaml
from docx.table import Table
from docx.text.paragraph import Paragraph


def iter_body_items(doc):
    """Yield each top-level body element as either a Paragraph or Table."""
    for child in doc.element.body:
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            yield Paragraph(child, doc)
        elif tag == 'tbl':
            yield Table(child, doc)


def cell_kv(cell):
    """Read a cell's nested 2-column key/value table into a dict.

    A value cell can itself hold a nested table (e.g. kodeliste-værdier for
    an attribut), which is expanded into a list of {værdi, definition}.
    """
    nested = cell.tables[0]
    kv = {}
    for row in nested.rows:
        key = row.cells[0].text.strip().rstrip(':')
        value_cell = row.cells[1]
        if value_cell.tables:
            kv[key] = [
                {'værdi': vrow.cells[0].text.strip(), 'definition': vrow.cells[1].text.strip()}
                for vrow in value_cell.tables[0].rows
            ]
        else:
            kv[key] = value_cell.text.strip()
    return kv


def parse_attribut(cell):
    return cell_kv(cell)


def parse_restriktion(cell):
    return cell_kv(cell)


def parse_associationsrolle(cell):
    return cell_kv(cell)


ROW_PARSERS = {
    'Attribut:': ('attributter', parse_attribut),
    'Restriktion:': ('restriktioner', parse_restriktion),
    'Associationsrolle': ('associationsroller', parse_associationsrolle),
}


def parse_featuretype(table, pakke):
    header_cell = table.rows[0].cells[0]
    featuretype = {
        'navn': header_cell.text.strip(),
        'pakke': pakke,
        **cell_kv(header_cell),
        'attributter': [],
        'restriktioner': [],
        'associationsroller': [],
    }
    for row in table.rows[1:]:
        label = row.cells[0].text.strip()
        entry = ROW_PARSERS.get(label)
        if entry is None:
            print(f'Ukendt rækketype {label!r} i {featuretype["navn"]!r}', file=sys.stderr)
            continue
        key, parser = entry
        featuretype[key].append(parser(row.cells[0]))
    return featuretype


def parse_featurekatalog(path):
    doc = docx.Document(path)
    featuretyper = []
    pakke = None
    for item in iter_body_items(doc):
        if isinstance(item, Paragraph) and item.style.name == 'Heading 1':
            pakke = item.text.strip()
        elif isinstance(item, Table):
            featuretyper.append(parse_featuretype(item, pakke))
    return featuretyper


EFTER_SKAERINGSDATO_RE = re.compile(r'efter\s+skæringsdatoen', re.IGNORECASE)


def all_restriktioner(featuretyper):
    """Group restriktioner by identical (Navn, Udtryk) and list which featuretyper share
    each exact rule. Sorted by Navn, then by group size descending, so the different
    variants of a reused Navn end up next to each other."""
    groups = {}
    for ft in featuretyper:
        for r in ft['restriktioner']:
            key = (r['Navn'], r['Udtryk'])
            group = groups.setdefault(key, {
                'Navn': r['Navn'],
                'Featuretyper': [],
                'Udtryk': r['Udtryk'],
                'EfterSkæringsdato': bool(EFTER_SKAERINGSDATO_RE.search(r['Udtryk'])),
            })
            group['Featuretyper'].append(ft['navn'])

    restriktioner = list(groups.values())
    restriktioner.sort(key=lambda r: (r['Navn'], -len(r['Featuretyper'])))
    return restriktioner


def slugify(navn):
    return re.sub(r'[^\w-]', '_', navn)


def _represent_str(dumper, data):
    """Dump multiline strings as YAML block-literals ('|-') instead of '\\n'-escaped."""
    style = '|' if '\n' in data else None
    return dumper.represent_scalar('tag:yaml.org,2002:str', data, style=style)


yaml.add_representer(str, _represent_str)
